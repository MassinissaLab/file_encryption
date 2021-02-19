
from RSA import *
import docx
import unicodedata
import codecs
filename="D:/WorkSpaces/Python workspace/tests/test.docx"

doc =docx.Document(filename)


d=docx.Document(filename)


ClearText =[]

for p in d.paragraphs :
	ClearText.append(p.text)
			
		

p = 11
q = 17
print ("Generating your public/private keypairs now . . .")

	
public, private = generer_cles(p, q)

print ("Your public key is ", public ," and your private key is ", private)



encrypted_msg = encrypt(public,ClearText[0])
x=''
for i in encrypted_msg:
	

	c=chr(i)
	
	x+=c


L=[]
L.append(x) 
 
fcript =codecs.open("ee.txt", "w","utf-8-sig")
fcript.writelines(L)
fcript.close()


K=[]
doc1 = docx.Document() 
with codecs.open("ee.txt", "r", "utf-8-sig") as myfile:
						
	Lines = myfile.readlines()
	
	for line in Lines:
		
		x=[]
		m=0
		for i in line:
			m = ord(i)
			x.append(m)

		

		decrypted_msg=decrypt(private,x)
		
		K.append(decrypted_msg)
		

fcript =codecs.open("eee.txt", "w","utf-8-sig")
fcript.writelines(K)
fcript.close()



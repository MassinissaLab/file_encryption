from PyQt5.QtGui import *
import os
from PyQt5.QtWidgets import *
from PyQt5.uic import loadUi
import time
import os,subprocess,sys
from PyQt5.QtGui import QIcon, QPixmap
from PyQt5.QtCore import Qt
from RSA import *
import codecs
import docx 
class Ui_cryptofichier(QMainWindow):
    global   rsapc,rsapr,FilePath,ResultPath
    def __init__(self):
        super(Ui_cryptofichier, self).__init__()
        loadUi("cryptofichier.ui", self)
        self.FilePath=""
        self.ResultPath=""
        self.msg_0.hide()
        self.msg_1.hide()
        self.msg_2.hide()
        self.msg_3.hide()
        self.msg_4.hide()
        self.msg_5.hide()
        self.msg_6.hide()
        self.msg_7.hide()
        
       
    
        self.confirmer_btn.clicked.connect(self.generercles)
        self.confirmer_btn_2.clicked.connect(self.generercles2)
        self.parcourire_btn.clicked.connect(self.choisirchemin)
        self.parcourire_btn_2.clicked.connect(self.rchoisirchemin)
        self.crypter_btn.clicked.connect(self.crypterf)
        self.decrypter_btn.clicked.connect(self.decrypterf)
          
    def generercles(self):
        self.msg_0.hide()
         
        
        try:
            public, private = generer_cles(self.rsap.value(), self.rsaq.value())
            self.rsapc=(public[0], public[1])
            self.rsapr=(private[0], private[1])
            self.clepub.setText(str(self.rsapc))
            self.cleprv.setText(str(self.rsapr))


        except :
              self.msg_0.show() 

    def generercles2(self):

            public, private = p_q_autocle()
            self.rsapc=(public[0], public[1])
            self.rsapr=(private[0], private[1])
            self.clepub.setText(str(self.rsapc))
            self.cleprv.setText(str(self.rsapr))

              

    def crypterf(self):
        self.FilePath=self.chemin.text()
        self.ResultPath=self.rchemin.text()
        self.msg_1.hide()
        self.msg_2.hide()
        self.msg_3.hide()
        self.msg_4.hide()
        self.msg_5.hide()
        self.msg_6.hide()
        self.msg_7.hide()

        L=[]
       
        print(self.rsapc)
        if self.FilePath=="":
            self.msg_1.show()
        elif self.ResultPath=="":
            self.msg_6.show()

        else:
            extension=os.path.splitext(self.FilePath)
            extr=os.path.splitext(self.ResultPath)
            if ".txt" not in extr[1]:
                self.msg_5.show()

            elif ".txt" in extension[1]:


                with open(self.FilePath,'r') as myfile:

                        Lines = myfile.readlines()
                    
                        for line in Lines:
                           
                            
                            encrypted_msg=encrypt(self.rsapc, line.strip())


                            x=''
                            for i in encrypted_msg:
                                
                     
                                c=chr(i)
                                
                                x+=c
                
                            
                            
                            L.append(x)

                        
                print("encrypted")
                print(L)
                if L==[]:
                    self.msg_2.show()
                else:    
                    
                    fcript =codecs.open(self.ResultPath, "w", "utf-8-sig")
                    fcript.writelines(L)
                    fcript .close()
                    self.msg_3.show()
            elif ".docx" in extension[1]:
                print("word docx")
                doc =docx.Document(self.FilePath)


                Text =[]
                
                for pr in doc.paragraphs :
                    Text.append(pr.text)

                
                if Text==[]:
                    self.msg_2.show()
                else:  
                    Q=[]
                    for p in Text:
                        encrypted_msg = encrypt(public,p)
                        x=''
                        for i in encrypted_msg:
                            

                            c=chr(i)
                            
                            x+=c

                        
                        
                        Q.append(x) 
                       
                    fcript =codecs.open(self.ResultPath, "w","utf-8-sig")
                    fcript.writelines(Q)
                    fcript.close()
                    self.msg_3.show()
        

    def decrypterf(self):
        
        self.FilePath=self.chemin.text()
        self.ResultPath=self.rchemin.text()
        self.msg_0.hide()
        self.msg_1.hide()
        self.msg_2.hide()
        self.msg_3.hide()
        self.msg_4.hide()
        self.msg_5.hide()
        self.msg_6.hide()
        self.msg_7.hide()

        #self.rsapr=(self.pk1.value(),self.pk2.value())
        print(self.rsapr)
        

        if self.FilePath=="":
            self.msg_1.show()
        elif self.ResultPath=="":
            self.msg_6.show()

        else:
            extension=os.path.splitext(self.FilePath)
            extr=os.path.splitext(self.ResultPath)

            if(".txt" not in extension[1]):
                self.msg_7.show()

            
            else :

                print("recuperer text")
                L=[]
                with codecs.open(self.FilePath, "r", "utf-8-sig") as myfile:
                                
                    Lines = myfile.readlines()
                    
                    for line in Lines:
                        
                        x=[]
                        m=0
                        for i in line:
                            m = ord(i)
                            x.append(m)

                        
                        
                        decrypted_msg=decrypt(self.rsapr,x)
                        
                        L.append(decrypted_msg)

                    
                if L==[]:

                    self.msg_2.show()

                else:  
                    if ".txt" in extr[1]:

                        fcript =codecs.open(self.ResultPath, "w","utf-8-sig")
                        fcript.writelines(L)
                        fcript.close()
                        self.msg_4.show()

                    elif ".docx" in extr[1]:


                        doc1 = docx.Document() 
                                
                        pp="".join(L)
                        
                        para = doc1.add_paragraph().add_run(pp) 

                        doc1.save(self.ResultPath)
                        self.msg_4.show()
                       

    
        
        
    def choisirchemin(self):
        self.msg_1.hide()
        self.msg_2.hide()
        self.msg_3.hide()
        self.msg_4.hide()
        self.msg_5.hide()
        self.msg_6.hide()
        self.msg_7.hide()
        self.FilePath,Filetype = QFileDialog.getOpenFileName(self)
        self.chemin.setText(self.FilePath)
    def rchoisirchemin(self):
        self.msg_1.hide()
        self.msg_2.hide()
        self.msg_3.hide()
        self.msg_4.hide()
        self.msg_5.hide()
        self.msg_6.hide()
        self.msg_7.hide()
        self.ResultPath,ftype = QFileDialog.getOpenFileName(self)
        self.rchemin.setText(self.ResultPath)
        
if __name__ == "__main__":
    app = QApplication(sys.argv)
    myapp = Ui_cryptofichier()
    myapp.show()
    sys.exit(app.exec_())

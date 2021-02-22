from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from PyQt5.QtCore import Qt
import os,subprocess,sys
from RSA import *
import docx 


class Ui_CRYPTOFICHIER(object):
    global   rsapc,rsapr,FilePath,ResultPath
    def setupUi(self, CRYPTOFICHIER):
        CRYPTOFICHIER.setObjectName("CRYPTOFICHIER")
        CRYPTOFICHIER.setWindowModality(QtCore.Qt.ApplicationModal)
        CRYPTOFICHIER.resize(680, 481)
        CRYPTOFICHIER.setStyleSheet("background-color: rgb(0, 52, 52);")
        self.centralwidget = QtWidgets.QWidget(CRYPTOFICHIER)
        self.centralwidget.setObjectName("centralwidget")
        self.layoutWidget = QtWidgets.QWidget(self.centralwidget)
        self.layoutWidget.setGeometry(QtCore.QRect(0, 0, 681, 30))
        self.layoutWidget.setObjectName("layoutWidget")
        self.horizontalLayout = QtWidgets.QHBoxLayout(self.layoutWidget)
        self.horizontalLayout.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.label_6 = QtWidgets.QLabel(self.layoutWidget)
        self.label_6.setText("")
        self.label_6.setPixmap(QtGui.QPixmap("C:/Users/ML/.designer/backup/image/32CYPHER.jpg"))
        self.label_6.setScaledContents(True)
        self.label_6.setObjectName("label_6")
        self.horizontalLayout.addWidget(self.label_6)
        spacerItem = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Minimum)
        self.horizontalLayout.addItem(spacerItem)
        self.label_35 = QtWidgets.QLabel(self.layoutWidget)
        self.label_35.setStyleSheet("color: rgb(0, 255, 0);\n"
"background:transparent;\n"
"font: 75 16pt \"Century Gothic\";\n"
"")
        self.label_35.setObjectName("label_35")
        self.horizontalLayout.addWidget(self.label_35)
        self.label_68 = QtWidgets.QLabel(self.layoutWidget)
        self.label_68.setStyleSheet("color: #ffffff;\n"
"background:transparent;\n"
"font: 75 16pt \"Century Gothic\";\n"
"")
        self.label_68.setObjectName("label_68")
        self.horizontalLayout.addWidget(self.label_68)
        spacerItem1 = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.MinimumExpanding, QtWidgets.QSizePolicy.Minimum)
        self.horizontalLayout.addItem(spacerItem1)
        self.decrypter_btn = QtWidgets.QPushButton(self.centralwidget)
        self.decrypter_btn.setEnabled(True)
        self.decrypter_btn.setGeometry(QtCore.QRect(370, 390, 161, 31))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.decrypter_btn.sizePolicy().hasHeightForWidth())
        self.decrypter_btn.setSizePolicy(sizePolicy)
        self.decrypter_btn.setStyleSheet("font: 14pt \"5by7\";\n"
"color: rgb(255, 255, 255);\n"
"background-color: rgb(6, 100, 135);\n"
"")
        self.decrypter_btn.setCheckable(False)
        self.decrypter_btn.setDefault(True)
        self.decrypter_btn.setObjectName("decrypter_btn")
        self.crypter_btn = QtWidgets.QPushButton(self.centralwidget)
        self.crypter_btn.setEnabled(True)
        self.crypter_btn.setGeometry(QtCore.QRect(180, 390, 161, 31))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.crypter_btn.sizePolicy().hasHeightForWidth())
        self.crypter_btn.setSizePolicy(sizePolicy)
        self.crypter_btn.setStyleSheet("background-color: rgb(6, 100, 135);\n"
"font: 14pt \"5by7\";\n"
"border-color: rgb(255, 255, 255);\n"
"color: rgb(255, 255, 255);font: 14pt \"5by7\";\n"
"")
        self.crypter_btn.setCheckable(False)
        self.crypter_btn.setDefault(True)
        self.crypter_btn.setObjectName("crypter_btn")
        self.msg_3 = QtWidgets.QLabel(self.centralwidget)
        self.msg_3.setEnabled(False)
        self.msg_3.setGeometry(QtCore.QRect(10, 440, 281, 31))
        self.msg_3.setStyleSheet("background: rgba(0, 170, 0,0.7);\n"
"\n"
"color: rgb(255, 255, 255);\n"
"border-radius:10px;\n"
"font: 12pt \"Century Gothic\";")
        self.msg_3.setObjectName("msg_3")
        self.msg_1 = QtWidgets.QLabel(self.centralwidget)
        self.msg_1.setEnabled(False)
        self.msg_1.setGeometry(QtCore.QRect(380, 220, 291, 31))
        self.msg_1.setStyleSheet("background:rgba(255,0,0,0.8);\n"
"color: rgb(255, 255, 255);\n"
"border-radius:10px;\n"
"font: 12pt \"Century Gothic\";")
        self.msg_1.setObjectName("msg_1")
        self.msg_2 = QtWidgets.QLabel(self.centralwidget)
        self.msg_2.setEnabled(False)
        self.msg_2.setGeometry(QtCore.QRect(10, 190, 301, 31))
        self.msg_2.setStyleSheet("background:rgba(255,0,0,0.8);\n"
"color: rgb(255, 255, 255);\n"
"border-radius:10px;\n"
"font: 12pt \"Century Gothic\";")
        self.msg_2.setObjectName("msg_2")
        self.msg_4 = QtWidgets.QLabel(self.centralwidget)
        self.msg_4.setEnabled(False)
        self.msg_4.setGeometry(QtCore.QRect(10, 440, 301, 31))
        self.msg_4.setStyleSheet("background: rgba(0, 170, 0,0.7);\n"
"\n"
"color: rgb(255, 255, 255);\n"
"border-radius:10px;\n"
"font: 12pt \"Century Gothic\";")
        self.msg_4.setObjectName("msg_4")
        self.chemin = QtWidgets.QLineEdit(self.centralwidget)
        self.chemin.setEnabled(True)
        self.chemin.setGeometry(QtCore.QRect(10, 260, 491, 31))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.chemin.sizePolicy().hasHeightForWidth())
        self.chemin.setSizePolicy(sizePolicy)
        self.chemin.setSizeIncrement(QtCore.QSize(15, 15))
        self.chemin.setBaseSize(QtCore.QSize(15, 15))
        self.chemin.setMouseTracking(False)
        self.chemin.setLayoutDirection(QtCore.Qt.LeftToRight)
        self.chemin.setAutoFillBackground(False)
        self.chemin.setStyleSheet("\n"
"\n"
"color:#ffffff;\n"
"border-radius: 0px;\n"
"border:1px solid #9cd3e2;\n"
"background-color: rgb(33, 65, 103);\n"
"font: 12pt \"Century Gothic\";")
        self.chemin.setInputMask("")
        self.chemin.setText("")
        self.chemin.setMaxLength(999999999)
        self.chemin.setFrame(True)
        self.chemin.setCursorPosition(0)
        self.chemin.setPlaceholderText("chemin de fichier à crypter / décrypter")
        self.chemin.setCursorMoveStyle(QtCore.Qt.LogicalMoveStyle)
        self.chemin.setClearButtonEnabled(False)
        self.chemin.setObjectName("chemin")
        self.parcourire_btn = QtWidgets.QPushButton(self.centralwidget)
        self.parcourire_btn.setEnabled(True)
        self.parcourire_btn.setGeometry(QtCore.QRect(510, 260, 161, 31))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.parcourire_btn.sizePolicy().hasHeightForWidth())
        self.parcourire_btn.setSizePolicy(sizePolicy)
        self.parcourire_btn.setStyleSheet("font: 14pt \"5by7\";\n"
"color: rgb(255, 255, 255);\n"
"background-color: rgb(6, 100, 135);\n"
"")
        self.parcourire_btn.setCheckable(True)
        self.parcourire_btn.setChecked(False)
        self.parcourire_btn.setAutoExclusive(False)
        self.parcourire_btn.setAutoDefault(False)
        self.parcourire_btn.setDefault(True)
        self.parcourire_btn.setFlat(False)
        self.parcourire_btn.setObjectName("parcourire_btn")
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setGeometry(QtCore.QRect(10, 70, 491, 31))
        self.label_3.setStyleSheet("background:transparent;\n"
"font: 14pt \"5by7\";\n"
"\n"
"font: 14pt \"Century Gothic\";\n"
"\n"
"color:#ffffff;")
        self.label_3.setObjectName("label_3")
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setGeometry(QtCore.QRect(10, 230, 361, 31))
        self.label_2.setStyleSheet("background:transparent;\n"
"font: 14pt \"5by7\";\n"
"\n"
"font: 14pt \"Century Gothic\";\n"
"\n"
"color:#ffffff;")
        self.label_2.setObjectName("label_2")
        self.confirmer_btn = QtWidgets.QPushButton(self.centralwidget)
        self.confirmer_btn.setGeometry(QtCore.QRect(510, 70, 161, 31))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.confirmer_btn.sizePolicy().hasHeightForWidth())
        self.confirmer_btn.setSizePolicy(sizePolicy)
        self.confirmer_btn.setStyleSheet("font: 14pt \"5by7\";\n"
"color: rgb(255, 255, 255);\n"
"background-color: rgb(6, 100, 135);\n"
"")
        self.confirmer_btn.setCheckable(False)
        self.confirmer_btn.setDefault(True)
        self.confirmer_btn.setObjectName("confirmer_btn")
        self.rchemin = QtWidgets.QLineEdit(self.centralwidget)
        self.rchemin.setEnabled(True)
        self.rchemin.setGeometry(QtCore.QRect(10, 340, 491, 31))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.rchemin.sizePolicy().hasHeightForWidth())
        self.rchemin.setSizePolicy(sizePolicy)
        self.rchemin.setSizeIncrement(QtCore.QSize(15, 15))
        self.rchemin.setBaseSize(QtCore.QSize(15, 15))
        self.rchemin.setMouseTracking(False)
        self.rchemin.setLayoutDirection(QtCore.Qt.LeftToRight)
        self.rchemin.setAutoFillBackground(False)
        self.rchemin.setStyleSheet("\n"
"\n"
"color:#ffffff;\n"
"border-radius: 0px;\n"
"border:1px solid #9cd3e2;\n"
"background-color: rgb(33, 65, 103);\n"
"font: 12pt \"Century Gothic\";")
        self.rchemin.setInputMask("")
        self.rchemin.setText("")
        self.rchemin.setMaxLength(999999999)
        self.rchemin.setFrame(True)
        self.rchemin.setCursorPosition(0)
        self.rchemin.setPlaceholderText("chemin de fichier resultant ")
        self.rchemin.setCursorMoveStyle(QtCore.Qt.LogicalMoveStyle)
        self.rchemin.setClearButtonEnabled(False)
        self.rchemin.setObjectName("rchemin")
        self.label_5 = QtWidgets.QLabel(self.centralwidget)
        self.label_5.setGeometry(QtCore.QRect(10, 310, 381, 31))
        self.label_5.setStyleSheet("background:transparent;\n"
"font: 14pt \"5by7\";\n"
"\n"
"font: 14pt \"Century Gothic\";\n"
"\n"
"color:#ffffff;")
        self.label_5.setObjectName("label_5")
        self.msg_5 = QtWidgets.QLabel(self.centralwidget)
        self.msg_5.setEnabled(False)
        self.msg_5.setGeometry(QtCore.QRect(400, 300, 261, 31))
        self.msg_5.setStyleSheet("background:rgba(255,0,0,0.8);\n"
"color: rgb(255, 255, 255);\n"
"border-radius:10px;\n"
"font: 12pt \"Century Gothic\";")
        self.msg_5.setObjectName("msg_5")
        self.parcourire_btn_2 = QtWidgets.QPushButton(self.centralwidget)
        self.parcourire_btn_2.setEnabled(True)
        self.parcourire_btn_2.setGeometry(QtCore.QRect(510, 340, 161, 31))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.parcourire_btn_2.sizePolicy().hasHeightForWidth())
        self.parcourire_btn_2.setSizePolicy(sizePolicy)
        self.parcourire_btn_2.setStyleSheet("background-color: rgb(6, 100, 135);\n"
"font: 14pt \"5by7\";\n"
"color: rgb(255, 255, 255);\n"
"")
        self.parcourire_btn_2.setCheckable(True)
        self.parcourire_btn_2.setDefault(True)
        self.parcourire_btn_2.setObjectName("parcourire_btn_2")
        self.prvkey = QtWidgets.QPushButton(self.centralwidget)
        self.prvkey.setGeometry(QtCore.QRect(510, 150, 161, 31))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.prvkey.sizePolicy().hasHeightForWidth())
        self.prvkey.setSizePolicy(sizePolicy)
        self.prvkey.setStyleSheet("font: 14pt \"5by7\";\n"
"color: rgb(255, 255, 255);\n"
"background-color: rgb(6, 100, 135);\n"
"")
        self.prvkey.setCheckable(True)
        self.prvkey.setDefault(True)
        self.prvkey.setObjectName("prvkey")
        self.pubkey = QtWidgets.QPushButton(self.centralwidget)
        self.pubkey.setGeometry(QtCore.QRect(510, 110, 161, 31))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.pubkey.sizePolicy().hasHeightForWidth())
        self.pubkey.setSizePolicy(sizePolicy)
        self.pubkey.setStyleSheet("font: 14pt \"5by7\";\n"
"color: rgb(255, 255, 255);\n"
"background-color: rgb(6, 100, 135);\n"
"")
        self.pubkey.setCheckable(True)
        self.pubkey.setDefault(True)
        self.pubkey.setObjectName("pubkey")
        self.chemin_prvkey = QtWidgets.QLineEdit(self.centralwidget)
        self.chemin_prvkey.setEnabled(True)
        self.chemin_prvkey.setGeometry(QtCore.QRect(10, 150, 491, 31))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.chemin_prvkey.sizePolicy().hasHeightForWidth())
        self.chemin_prvkey.setSizePolicy(sizePolicy)
        self.chemin_prvkey.setSizeIncrement(QtCore.QSize(15, 15))
        self.chemin_prvkey.setBaseSize(QtCore.QSize(15, 15))
        self.chemin_prvkey.setMouseTracking(False)
        self.chemin_prvkey.setLayoutDirection(QtCore.Qt.LeftToRight)
        self.chemin_prvkey.setAutoFillBackground(False)
        self.chemin_prvkey.setStyleSheet("\n"
"\n"
"color:#ffffff;\n"
"border-radius: 0px;\n"
"border:1px solid #9cd3e2;\n"
"background-color: rgb(33, 65, 103);\n"
"font: 12pt \"Century Gothic\";")
        self.chemin_prvkey.setInputMask("")
        self.chemin_prvkey.setText("")
        self.chemin_prvkey.setMaxLength(999999999)
        self.chemin_prvkey.setFrame(True)
        self.chemin_prvkey.setCursorPosition(0)
        self.chemin_prvkey.setPlaceholderText("chemin de la clé privé  .pem")
        self.chemin_prvkey.setCursorMoveStyle(QtCore.Qt.LogicalMoveStyle)
        self.chemin_prvkey.setClearButtonEnabled(False)
        self.chemin_prvkey.setObjectName("chemin_prvkey")
        self.chemin_pubkey = QtWidgets.QLineEdit(self.centralwidget)
        self.chemin_pubkey.setEnabled(True)
        self.chemin_pubkey.setGeometry(QtCore.QRect(10, 110, 491, 31))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.chemin_pubkey.sizePolicy().hasHeightForWidth())
        self.chemin_pubkey.setSizePolicy(sizePolicy)
        self.chemin_pubkey.setSizeIncrement(QtCore.QSize(15, 15))
        self.chemin_pubkey.setBaseSize(QtCore.QSize(15, 15))
        self.chemin_pubkey.setMouseTracking(False)
        self.chemin_pubkey.setLayoutDirection(QtCore.Qt.LeftToRight)
        self.chemin_pubkey.setAutoFillBackground(False)
        self.chemin_pubkey.setStyleSheet("\n"
"\n"
"color:#ffffff;\n"
"border-radius: 0px;\n"
"border:1px solid #9cd3e2;\n"
"background-color: rgb(33, 65, 103);\n"
"font: 12pt \"Century Gothic\";")
        self.chemin_pubkey.setInputMask("")
        self.chemin_pubkey.setText("")
        self.chemin_pubkey.setMaxLength(999999999)
        self.chemin_pubkey.setFrame(True)
        self.chemin_pubkey.setCursorPosition(0)
        self.chemin_pubkey.setPlaceholderText("chemin de la clé public  .pem")
        self.chemin_pubkey.setCursorMoveStyle(QtCore.Qt.LogicalMoveStyle)
        self.chemin_pubkey.setClearButtonEnabled(False)
        self.chemin_pubkey.setObjectName("chemin_pubkey")
        self.msg_6 = QtWidgets.QLabel(self.centralwidget)
        self.msg_6.setEnabled(False)
        self.msg_6.setGeometry(QtCore.QRect(360, 190, 311, 31))
        self.msg_6.setStyleSheet("background:rgba(255,0,0,0.8);\n"
"color: rgb(255, 255, 255);\n"
"border-radius:10px;\n"
"font: 12pt \"Century Gothic\";")
        self.msg_6.setObjectName("msg_6")
        self.msg_7 = QtWidgets.QLabel(self.centralwidget)
        self.msg_7.setEnabled(False)
        self.msg_7.setGeometry(QtCore.QRect(240, 40, 191, 31))
        self.msg_7.setStyleSheet("background: rgba(0, 170, 0,0.7);\n"
"\n"
"color: rgb(255, 255, 255);\n"
"border-radius:10px;\n"
"font: 12pt \"Century Gothic\";")
        self.msg_7.setObjectName("msg_7")
        self.msg_8 = QtWidgets.QLabel(self.centralwidget)
        self.msg_8.setEnabled(False)
        self.msg_8.setGeometry(QtCore.QRect(360, 190, 311, 31))
        self.msg_8.setStyleSheet("background:rgba(255,0,0,0.8);\n"
"color: rgb(255, 255, 255);\n"
"border-radius:10px;\n"
"font: 12pt \"Century Gothic\";")
        self.msg_8.setObjectName("msg_8")
        self.line = QtWidgets.QFrame(self.centralwidget)
        self.line.setGeometry(QtCore.QRect(0, 30, 681, 3))
        self.line.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.line.setFrameShape(QtWidgets.QFrame.HLine)
        self.line.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.line.setObjectName("line")
        CRYPTOFICHIER.setCentralWidget(self.centralwidget)

        self.retranslateUi(CRYPTOFICHIER)
        QtCore.QMetaObject.connectSlotsByName(CRYPTOFICHIER)

    def retranslateUi(self, CRYPTOFICHIER):
        _translate = QtCore.QCoreApplication.translate
        CRYPTOFICHIER.setWindowTitle(_translate("CRYPTOFICHIER", "CRYPTOFICHIER"))
        self.label_35.setText(_translate("CRYPTOFICHIER", "Cryptage "))
        self.label_68.setText(_translate("CRYPTOFICHIER", "Fichier Word .docx"))
        self.decrypter_btn.setText(_translate("CRYPTOFICHIER", "Decrypter"))
        self.decrypter_btn.setShortcut(_translate("CRYPTOFICHIER", "Return"))
        self.crypter_btn.setText(_translate("CRYPTOFICHIER", "Crypter"))
        self.crypter_btn.setShortcut(_translate("CRYPTOFICHIER", "Return"))
        self.msg_3.setText(_translate("CRYPTOFICHIER", " Le fichier a été crypté avec succes"))
        self.msg_1.setText(_translate("CRYPTOFICHIER", " Entrer le chemin d\'abord "))
        self.msg_2.setText(_translate("CRYPTOFICHIER", " le fichier n\'existe pas ou type incorrect"))
        self.msg_4.setText(_translate("CRYPTOFICHIER", " Le fichier a été decrypté avec succes"))
        self.parcourire_btn.setText(_translate("CRYPTOFICHIER", "Parcourire"))
        self.parcourire_btn.setShortcut(_translate("CRYPTOFICHIER", "Return"))
        self.label_3.setText(_translate("CRYPTOFICHIER", " Auto généreration  des es fichier .pem des cléfs RSA"))
        self.label_2.setText(_translate("CRYPTOFICHIER", "Chemin Du fichier à crypter/ decrypter"))
        self.confirmer_btn.setText(_translate("CRYPTOFICHIER", "Au-Gen Cléfs"))
        self.confirmer_btn.setShortcut(_translate("CRYPTOFICHIER", "Return"))
        self.label_5.setText(_translate("CRYPTOFICHIER", "Chemin Du fichier resultant "))
        self.msg_5.setText(_translate("CRYPTOFICHIER", "Indiquer chemin fichier resultant "))
        self.parcourire_btn_2.setText(_translate("CRYPTOFICHIER", "Parcourire"))
        self.parcourire_btn_2.setShortcut(_translate("CRYPTOFICHIER", "Return"))
        self.prvkey.setText(_translate("CRYPTOFICHIER", "Fichier Clef privé"))
        self.prvkey.setShortcut(_translate("CRYPTOFICHIER", "Return"))
        self.pubkey.setText(_translate("CRYPTOFICHIER", "Fichier Clef public"))
        self.pubkey.setShortcut(_translate("CRYPTOFICHIER", "Return"))
        self.msg_6.setText(_translate("CRYPTOFICHIER", "le fichier des clef doit être /fichier.pem"))
        self.msg_7.setText(_translate("CRYPTOFICHIER", "Les Cléfs sont génerées"))
        self.msg_8.setText(_translate("CRYPTOFICHIER", "Votre cléf privée est incorrecte"))
        self.FilePath=""
        self.ResultPath=""
        
        self.msg_1.hide()
        self.msg_2.hide()
        self.msg_3.hide()
        self.msg_4.hide()
        self.msg_5.hide()
        self.msg_6.hide()
        self.msg_7.hide()
        self.msg_8.hide()
       
    
        self.confirmer_btn.clicked.connect(self.generercles)

        self.pubkey.clicked.connect(self.choisirchemin)
        self.prvkey.clicked.connect(self.choisirchemin)
        self.parcourire_btn.clicked.connect(self.choisirchemin)
        self.parcourire_btn_2.clicked.connect(self.choisirchemin)

        self.crypter_btn.clicked.connect(self.crypterf)
        self.decrypter_btn.clicked.connect(self.decrypterf)

    def generercles(self):
        
        public, private =newkeys(1024)
        self.msg_7.show()


              

    def crypterf(self):
        self.msg_1.hide()
        self.msg_2.hide()
        self.msg_3.hide()
        self.msg_4.hide()
        self.msg_5.hide()
        self.msg_6.hide()
        self.msg_7.hide()
        self.msg_8.hide()

        self.FilePath=self.chemin.text()
        self.ResultPath=self.rchemin.text()


        ext1=os.path.splitext(self.FilePath)
        ext2=os.path.splitext(self.ResultPath)


        if self.FilePath=="":
            self.msg_1.show()
        elif self.ResultPath=="":
            self.msg_5.show()
        else:
            if(".docx" not in ext1[1] or ".docx" not in ext2[1]) :
                self.msg_2.show()

            else :
                print("encrypt")
                print(self.rsapc)
                doc =docx.Document(self.FilePath)
                ClearText =[]
                cc=[]

                for p in doc.paragraphs :
                    ClearText.append(p.text)

                for el in ClearText:
                    
                    cc.append(el.split(" "))

                

                
                docc = docx.Document() 
                        


                result=[]
                
                for p in cc:
                    pf=[]
                    for e in p :
                        
                        pf.extend(str(base64.b64encode(encrypt(e.encode('utf-8'),self.rsapc)), encoding='utf-8')+" ")

                    result.append(pf)

                print(result)
                for pp in result :
                    para = docc.add_paragraph().add_run(pp)
                docc.save(self.ResultPath)

                self.msg_3.show()



    def decrypterf(self):
        self.msg_1.hide()
        self.msg_2.hide()
        self.msg_3.hide()
        self.msg_4.hide()
        self.msg_5.hide()
        self.msg_6.hide()
        self.msg_7.hide()
        self.msg_8.hide()

        self.FilePath=self.chemin.text()
        self.ResultPath=self.rchemin.text()
        

        extension=os.path.splitext(self.FilePath)
        extr=os.path.splitext(self.ResultPath)

        if self.FilePath=="":
            self.msg_1.show()
        elif self.ResultPath=="":
            self.msg_5.show()
        else:
            if(".docx" not in extension[1] or ".docx" not in extr[1]):
                self.msg_2.show()
            
            else :
                docc =docx.Document(self.FilePath)
                print(self.FilePath)
                

                CypherText =[]
                ccc=[]
                for p in docc.paragraphs :
                    CypherText.append(p.text)


                for el in CypherText:
                    ccc.append(el.split(" "))

                dresult=[]
                for pr in ccc:
                    pf=[]
                    for  ll in pr :
                        tmp=base64.b64decode(ll)
                        
                        if tmp!=b'':

                            try:
                                pf.append(decrypt(tmp,self.rsapr).decode('utf-8')+" ")
                            except:
                                self.msg_8.show()
                            

                    dresult.append(pf)


                ddocc =docx.Document()
                for ppp in dresult :
                    par = ddocc.add_paragraph().add_run(ppp)
                ddocc.save(self.ResultPath)


                self.msg_4.show()
        
    def choisirchemin(self):
        self.msg_1.hide()
        self.msg_2.hide()
        self.msg_3.hide()
        self.msg_4.hide()
        self.msg_5.hide()
        self.msg_6.hide()
        self.msg_7.hide()
        self.msg_8.hide()
        


        if self.parcourire_btn.isChecked():
            FPath,Ftype = QFileDialog.getOpenFileName()
            self.parcourire_btn.setChecked(False)
            self.chemin.setText(FPath)


        elif self.parcourire_btn_2.isChecked():
            FPath,Ftype = QFileDialog.getOpenFileName()
            self.parcourire_btn_2.setChecked(False)
            self.rchemin.setText(FPath)
           


        elif self.pubkey.isChecked():
           FPath,Ftype = QFileDialog.getOpenFileName()
           ext=os.path.splitext(FPath)
           self.pubkey.setChecked(False)
           self.chemin_pubkey.setText(FPath)
           if self.chemin_pubkey.text()=="" or ".pem" not in ext[1]:
                self.msg_6.show()
           else:
            self.rsapc=importKey(self.chemin_pubkey.text())


        elif self.prvkey.isChecked():
            FPath,Ftype = QFileDialog.getOpenFileName()
            self.prvkey.setChecked(False)
            ext=os.path.splitext(FPath)
            self.chemin_prvkey.setText(FPath)
            if self.chemin_prvkey.text()=="":
                self.msg_6.show()
            else:
                self.rsapr=importKey(self.chemin_prvkey.text())


if __name__ == "__main__":
    import sys
    from os import environ
    environ["QT_DEVICE_PIXEL_RATIO"] = "0"
    environ["QT_AUTO_SCREEN_SCALE_FACTOR"] = "1"
    environ["QT_SCREEN_SCALE_FACTORS"] = "1"
    environ["QT_SCALE_FACTOR"] = "1"

    app = QtWidgets.QApplication(sys.argv)
    CRYPTOFICHIER = QtWidgets.QMainWindow()
    ui = Ui_CRYPTOFICHIER()
    ui.setupUi(CRYPTOFICHIER)
    
    CRYPTOFICHIER.show()
    sys.exit(app.exec_())

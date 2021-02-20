from PyQt5.QtGui import *
import os
from PyQt5.QtWidgets import *
from PyQt5.uic import loadUi
import time
import os,subprocess,sys
from PyQt5.QtGui import QIcon, QPixmap
from PyQt5.QtCore import Qt
from RSA import *

import docx 
class Ui_cryptofichier(QMainWindow):
    global   rsapc,rsapr,FilePath,ResultPath
    def __init__(self):
        super(Ui_cryptofichier, self).__init__()
        loadUi("cryptofichier.ui", self)
        self.FilePath=""
        self.ResultPath=""
        
        self.msg_1.hide()
        self.msg_2.hide()
        self.msg_3.hide()
        self.msg_4.hide()
        self.msg_5.hide()
        self.msg_6.hide()
        self.msg_7.hide()
       
    
        self.confirmer_btn.clicked.connect(self.generercles)

        self.pubkey.clicked.connect(self.choisirchemin)
        self.prvkey.clicked.connect(self.choisirchemin)
        self.parcourire_btn.clicked.connect(self.choisirchemin)
        self.parcourire_btn_2.clicked.connect(self.choisirchemin)

        self.crypter_btn.clicked.connect(self.crypterf)
        self.decrypter_btn.clicked.connect(self.decrypterf)
          
    def generercles(self):
        
        public, private =newkeys(1024)
        self.msg_7.show()


              

    def crypterf(self):
        self.msg_1.hide()
        self.msg_2.hide()
        self.msg_3.hide()
        self.msg_4.hide()
        self.msg_5.hide()
        self.msg_6.hide()
        self.msg_7.hide()

        self.FilePath=self.chemin.text()
        self.ResultPath=self.rchemin.text()


        ext1=os.path.splitext(self.FilePath)
        ext2=os.path.splitext(self.ResultPath)


        if self.FilePath=="":
            self.msg_1.show()
        elif self.ResultPath=="":
            self.msg_5.show()
        else:
            if(".docx" not in ext1[1] or ".docx" not in ext2[1]) :
                self.msg_2.show()

            else :
                print("encrypt")
                print(self.rsapc)
                doc =docx.Document(self.FilePath)
                ClearText =[]
                cc=[]

                for p in doc.paragraphs :
                    ClearText.append(p.text)

                for el in ClearText:
                    
                    cc.append(el.split(" "))

                

                
                docc = docx.Document() 
                        


                result=[]
                
                for p in cc:
                    pf=[]
                    for e in p :
                        
                        pf.extend(str(base64.b64encode(encrypt(e.encode('utf-8'),self.rsapc)), encoding='utf-8')+" ")

                    result.append(pf)

                print(result)
                for pp in result :
                    para = docc.add_paragraph().add_run(pp)
                docc.save(self.ResultPath)

                self.msg_3.show()



    def decrypterf(self):
        self.msg_1.hide()
        self.msg_2.hide()
        self.msg_3.hide()
        self.msg_4.hide()
        self.msg_5.hide()
        self.msg_6.hide()
        self.msg_7.hide()

        self.FilePath=self.chemin.text()
        self.ResultPath=self.rchemin.text()
        

        extension=os.path.splitext(self.FilePath)
        extr=os.path.splitext(self.ResultPath)

        if self.FilePath=="":
            self.msg_1.show()
        elif self.ResultPath=="":
            self.msg_5.show()
        else:
            if(".docx" not in extension[1] or ".docx" not in extr[1]):
                self.msg_2.show()
            
            else :
                docc =docx.Document(self.FilePath)
                print(self.FilePath)
                

                CypherText =[]
                ccc=[]
                for p in docc.paragraphs :
                    CypherText.append(p.text)


                for el in CypherText:
                    ccc.append(el.split(" "))

                dresult=[]
                for pr in ccc:
                    pf=[]
                    for  ll in pr :
                        tmp=base64.b64decode(ll)
                        
                        if tmp!=b'':


                            pf.append(decrypt(tmp,self.rsapr).decode('utf-8')+" ")
                            

                    dresult.append(pf)


                ddocc =docx.Document()
                for ppp in dresult :
                    par = ddocc.add_paragraph().add_run(ppp)
                ddocc.save(self.ResultPath)


                self.msg_4.show()
        
    def choisirchemin(self):
        self.msg_1.hide()
        self.msg_2.hide()
        self.msg_3.hide()
        self.msg_4.hide()
        self.msg_5.hide()
        self.msg_6.hide()
        


        if self.parcourire_btn.isChecked():
            FPath,Ftype = QFileDialog.getOpenFileName(self)
            self.chemin.setText(FPath)


        elif self.parcourire_btn_2.isChecked():
            FPath,Ftype = QFileDialog.getOpenFileName(self)
            self.rchemin.setText(FPath)
           


        elif self.pubkey.isChecked():
           FPath,Ftype = QFileDialog.getOpenFileName(self)
           ext=os.path.splitext(FPath)
           self.chemin_pubkey.setText(FPath)
           if self.chemin_pubkey.text()=="" or ".pem" not in ext[1]:
                self.msg_6.show()
           else:
            self.rsapc=importKey(self.chemin_pubkey.text())


        elif self.prvkey.isChecked():
            FPath,Ftype = QFileDialog.getOpenFileName(self)
            ext=os.path.splitext(FPath)
            self.chemin_prvkey.setText(FPath)
            if self.chemin_prvkey.text()=="":
                self.msg_6.show()
            else:
                self.rsapr=importKey(self.chemin_prvkey.text())



        
if __name__ == "__main__":
    app = QApplication(sys.argv)
    myapp = Ui_cryptofichier()
    myapp.show()
    sys.exit(app.exec_())
